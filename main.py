import os
import io
import time
import google.generativeai as genai
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
import fitz  # PyMuPDF
import logging
import streamlit as st
from reportlab.lib.pagesizes import letter, A4
from reportlab.pdfgen import canvas
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak, Frame, PageTemplate
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.enums import TA_JUSTIFY, TA_LEFT, TA_CENTER, TA_RIGHT
from google.api_core import exceptions as google_api_exceptions
from PIL import Image
import zipfile
from tqdm import tqdm
import contextlib  


# --- Logging Setup ---
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

# --- Helper Functions ---

 @contextlib.contextmanager
 def st_tqdm(iterable, desc=None, total=None, unit="it", **tqdm_kwargs):
     """Context manager for tqdm progress bar in Streamlit."""
     placeholder = st.empty()
     with tqdm(iterable, desc=desc, total=total, unit=unit, **tqdm_kwargs) as pbar:
         for item in pbar:
             yield item
             placeholder.write(pbar)


def create_header(document, text):
    """Adds a header to each page of a docx document."""
    for section in document.sections:
        header = section.header
        paragraph = header.paragraphs[0]
        paragraph.text = text
        paragraph_format = paragraph.paragraph_format
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

def extract_text_from_docx(docx_bytes):
    """Extracts text from a .docx file."""
    try:
        doc = Document(io.BytesIO(docx_bytes))
        return "\n\n".join([paragraph.text for paragraph in doc.paragraphs])
    except Exception as e:
        logging.error(f"Error extracting text from DOCX: {e}")
        st.error(f"Error extracting text from DOCX: {e}")
        return ""

def extract_text_from_pdf(pdf_bytes):
    """Extracts text from a .pdf file using PyMuPDF."""
    text = ""
    try:
        with fitz.open(stream=pdf_bytes, filetype="pdf") as doc:
            text = "".join([page.get_text() for page in doc])
        return text
    except Exception as e:
        logging.error(f"Error extracting text from PDF: {e}")
        st.error(f"Error extracting text from PDF: {e}")
        return ""

def split_into_paragraphs(text):
    """Splits text into paragraphs based on double newlines."""
    return [p.strip() for p in text.split("\n\n") if p.strip()]

def create_translation_prompt(source_language, target_language, document_summary, paragraph):
    """Creates the translation prompt for Gemini."""
    return f"""You are a professional translator. Translate the following paragraph from {source_language} to {target_language}.
Maintain the original meaning and tone as closely as possible.  Be as accurate as possible.

Here is a summary of the entire document for context:
{document_summary}

Paragraph to translate:
{paragraph}"""

def translate_paragraph(paragraph, source_language, target_language, document_summary, retries=3):
    """Translates a single paragraph using Gemini, with retries."""
    prompt = create_translation_prompt(source_language, target_language, document_summary, paragraph)
    for attempt in range(retries):
        try:
            start_time = time.time()  # Time the API call
            response = model.generate_content(prompt)
            end_time = time.time()  # Time the API call
            api_call_time = end_time - start_time  # Calculate the actual API call time

            if response.text:
                return response.text, "translated", api_call_time  # Return actual API call time
            else:
                logging.warning(f"Empty response from Gemini on attempt {attempt + 1}.")
                st.warning(f"Empty response from Gemini on attempt {attempt + 1}.")
                return "", "failed", 0  # Return 0 for failed attempts

        except google_api_exceptions.ClientError as e:
            logging.error(f"Gemini API error on attempt {attempt + 1}: {e}")
            st.error(f"Gemini API error on attempt {attempt + 1}: {e}")
            if e.code == 400 and "API key not valid" in str(e):
                raise ValueError("Invalid API Key provided.") from e
            #  No special handling for 429 errors, just retry
            elif attempt < retries - 1:
                st.info(f"Retrying...")  # Simplified retry message
                logging.info(f"Retrying...")
            else:
                return "", "failed", 0  # Return 0 delay on failure

        except Exception as e:
            logging.error(f"Unexpected error on attempt {attempt + 1}: {e}")
            st.error(f"Unexpected error on attempt {attempt + 1}: {e}")
            if attempt < retries - 1:
                st.info(f"Retrying...")
                logging.info(f"Retrying...")
            else:
                return "", "failed", 0  # Return 0 delay on failure

    return "", "failed", 0  # Return 0 delay if all retries fail


def generate_summary(text, target_language, max_length=700):
    """Generates a summary of the document using Gemini in the target language."""
    prompt = f"""Summarize the following text in {target_language} (not in markdown) in no more than {max_length} characters:\n\n{text}"""

    try:
        response = model.generate_content(prompt)
        return response.text if response.text else "Summary generation failed."
    except Exception as e:
        logging.error(f"Error generating summary: {e}")
        st.error(f"Error generating summary: {e}")
        return "Summary generation failed."

def set_paragraph_rtl(paragraph):
    """Sets the paragraph direction to Right-to-Left."""
    pPr = paragraph._element.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    pPr.append(bidi)
    paragraph_format = paragraph.paragraph_format
    paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

def set_table_rtl(table):
    """Sets the table direction to Right-to-Left."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.append(tblPr)
    bidi_visual = OxmlElement('w:bidiVisual')
    tblPr.append(bidi_visual)

def create_pdf_from_paragraphs(paragraphs, filename, is_rtl=False):
    """Creates a PDF from a list of paragraphs."""
    doc = SimpleDocTemplate(filename, pagesize=A4)
    elements = []

    styles = getSampleStyleSheet()
    style = styles["Normal"]
    style.alignment = TA_JUSTIFY
    if is_rtl:
        style.alignment = TA_RIGHT
        style.firstLineIndent = 0
        style.rightIndent = 0

    for para_text in paragraphs:
        p = Paragraph(para_text, style)
        elements.append(p)
        available_height = doc.height - doc.bottomMargin - doc.topMargin
        if elements:
            y = elements[-1].getSpaceAfter()
            available_height -= y
        w, h = p.wrap(doc.width, available_height)
        if h > available_height:
            elements.append(PageBreak())

    def header_footer(canvas, doc):
        canvas.saveState()
        styles = getSampleStyleSheet()
        header = Paragraph("Translated with AI, by Yedidya Harris", styles['Normal'])
        header.wrapOn(canvas, doc.width, doc.topMargin)
        header.drawOn(canvas, doc.leftMargin, doc.height + doc.topMargin - header.height)
        canvas.restoreState()

    page_template = PageTemplate(id='basic', frames=[Frame(doc.leftMargin, doc.bottomMargin, doc.width, doc.height)], onPage=header_footer)
    doc.addPageTemplates([page_template])
    doc.build(elements)

# --- Configuration and Model Setup ---
def configure_gemini(api_key):
    """Configures the Gemini AI model."""
    genai.configure(api_key=api_key)
    
    generation_config = {
        "temperature": 0.1,  # More deterministic, less "creative"
        "top_p": 0.95,
        "top_k": 64,
        "max_output_tokens": 8192,
    }
    
    model = genai.GenerativeModel(
        model_name="gemini-2.0-flash",
        generation_config=generation_config,
    )
    return model

# --- Streamlit UI ---

def main():
    """Main function to run the Streamlit translation application."""

    st.set_page_config(page_title="AI Document Translator", page_icon=":globe_with_meridians:", layout="wide")

    # --- Sidebar for Instructions and Information ---
    with st.sidebar:
        st.title("About")
        api_key_input = st.text_input("Enter your Gemini API Key (optional)", type="password")

        st.markdown(
            "This app translates documents using Google's Gemini (2.0 flash) model.  It supports `.docx` and `.pdf` files."
            " Upload your file, select the source and target languages, and click 'Translate'."
            " The app provides both a combined (source/target) and a translated-only version of your document."
        )
        st.markdown("---")
        st.markdown("**Developed by Yedidya Harris**")
        st.markdown(
            "This is a demonstration project and may have limitations.  For critical translations, always consult a professional human translator."
        )
        st.markdown("---")
        st.markdown("**Language Support:**")
        st.markdown("The app supports a wide range of languages.  If you encounter any issues with a specific language, please let me know.")

        # Add a nice image to the sidebar
        try:
            image = Image.open('translator_image.png')  # Replace with your image path
            st.image(image, caption='AI Translation', use_container_width=True)
        except FileNotFoundError:
            st.warning("Image file not found.  Please add translator_image.png to your project.")


    # --- Configuration and Model Setup ---
    # Use Streamlit secrets for API key management
    # --- API Key Logic ---
    if api_key_input:
        api_key_to_use = api_key_input
        st.sidebar.success("Using entered API key.")  # Optional: feedback
    else:
        api_key_to_use = st.secrets["GEMINI_API_KEY"]
        st.sidebar.info("Using API key from secrets.")  # Optional: feedback
    
    # **Initialize the Gemini model globally**
    global model  
    model = configure_gemini(api_key_to_use)
    

    
    # --- Main Content Area ---
    st.title("AI Document Translator")
    st.write("Upload a .docx or .pdf file to begin.")

    # --- Placeholders OUTSIDE of any columns or spinners ---
    progress_bar = st.progress(0)  # Initialize progress bar
    eta_placeholder = st.empty()  # Placeholder for ETA display

    uploaded_file = st.file_uploader("Choose a file", type=["docx", "pdf"])

    if uploaded_file is not None:
        file_content = uploaded_file.read()
        filename = uploaded_file.name

        language_options = [
            ('English', 'en'), ('Spanish', 'es'), ('French', 'fr'), ('German', 'de'),
            ('Chinese (Simplified)', 'zh-CN'), ('Chinese (Traditional)', 'zh-TW'),
            ('Japanese', 'ja'), ('Korean', 'ko'), ('Russian', 'ru'),
            ('Arabic', 'ar'), ('Hebrew', 'he'), ('Portuguese', 'pt'),
            ('Italian', 'it'), ('Dutch', 'nl'), ('Swedish', 'sv'),
            ('Norwegian', 'no'), ('Danish', 'da'), ('Finnish', 'fi'),
            ('Turkish', 'tr'), ('Indonesian', 'id'), ('Vietnamese', 'vi'),
            ('Greek', 'el'), ('Polish', 'pl'), ('Czech', 'cs'),
            ('Hungarian', 'hu'), ('Romanian', 'ro'), ('Thai', 'th'),
            ('Hindi', 'hi')
        ]

        col1, col2 = st.columns(2)
        with col1:
            source_language_tuple = st.selectbox("Source Language", options=language_options, format_func=lambda x: x[0], key="source_lang")
            source_language_name = source_language_tuple[0]
            source_language_code = source_language_tuple[1]

        with col2:
            target_language_tuple = st.selectbox("Target Language", options=language_options, format_func=lambda x: x[0], key="target_lang")
            target_language_name = target_language_tuple[0]
            target_language_code = target_language_tuple[1]


        if st.button("Translate"):
            if source_language_name == target_language_name:
                st.error("Source and target languages cannot be the same.")
                return

            is_target_rtl = target_language_code.lower() in ['he', 'ar', 'fa', 'ur', 'yi']

            with st.spinner("Processing document..."):
                if filename.endswith(".docx"):
                    text = extract_text_from_docx(file_content)
                elif filename.endswith(".pdf"):
                    text = extract_text_from_pdf(file_content)
                else:
                    st.error("Unsupported file type.")
                    return

                if not text:
                    st.error("Could not extract text from the document.")
                    return

                paragraphs = split_into_paragraphs(text)
                num_paragraphs = len(paragraphs)
                try:
                    document_summary = generate_summary(text, target_language_name)
                    st.success(f"Document summary generated in {target_language_name}.")
                    with st.expander("Show Summary in Target Language"):
                        st.write(document_summary)
                except ValueError as e:
                    st.error(f"Error: {e}")
                    return
                except Exception as e:
                    st.error(f"Error generating summary: {e}")
                    return

            with st.spinner("Translating..."):
                df_data = []
                translated_paragraphs = []
                start_time = time.time()
                total_api_time = 0

             with st_tqdm(paragraphs, desc="Translating Paragraphs", unit="paragraph") as pbar:
                 for i, paragraph in enumerate(pbar):
                     try:
                        translated_text, status, api_call_time = translate_paragraph(paragraph, source_language_name, target_language_name, document_summary)
                        total_api_time += api_call_time

                        df_data.append({
                            "paragraph_id": i + 1,
                            "source_text": paragraph,
                            "target_text": translated_text,
                            "status": status
                        })
                        if status == "translated":
                            translated_paragraphs.append(translated_text)
                    except ValueError as e:
                        st.error(f"Error: {e}")
                        return

                    # --- Progress Bar and ETA Calculation ---
                    progress = (i + 1) / num_paragraphs
                    progress_bar.progress(progress)

                    if i > 0:
                        elapsed_time = time.time() - start_time
                        estimated_total_time = (total_api_time / progress) + (10 * (num_paragraphs - i -1)) # Add remaining delay
                        remaining_time = estimated_total_time - elapsed_time
                        eta_placeholder.write(f"Estimated time remaining: {remaining_time:.2f} seconds")

                    time.sleep(10)  # 10-second delay after EACH paragraph


                df = pd.DataFrame(df_data)
                st.success("Translation complete!")

            # --- Display Results ---
            st.subheader("Translation Results")
            with st.expander("Show Full Translation Table"):
                st.dataframe(df)

            # --- DOCX Output ---
            with st.spinner("Generating DOCX files..."):
                combined_doc = Document()
                create_header(combined_doc, "Translated with AI, by Yedidya Harris")
                table = combined_doc.add_table(rows=1, cols=3)
                table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Paragraph ID'
                hdr_cells[1].text = 'Source Text'
                hdr_cells[2].text = 'Target Text'

                if is_target_rtl:
                    set_table_rtl(table)

                for index, row in df.iterrows():
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(row['paragraph_id'])
                    row_cells[1].text = row['source_text']
                    row_cells[2].text = row['target_text']
                    if is_target_rtl:
                        for paragraph in row_cells[2].paragraphs:
                            set_paragraph_rtl(paragraph)
                            paragraph_format = paragraph.paragraph_format
                            paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    else:
                        for paragraph in row_cells[2].paragraphs:
                            paragraph_format = paragraph.paragraph_format
                            paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

                combined_doc_filename = "combined_translation.docx"
                combined_doc.save(combined_doc_filename)

                translated_doc = Document()
                create_header(translated_doc, "Translated with AI, by Yedidya Harris")
                for paragraph_text in translated_paragraphs:
                    paragraph = translated_doc.add_paragraph(paragraph_text)
                    if is_target_rtl:
                        set_paragraph_rtl(paragraph)
                    paragraph_format = paragraph.paragraph_format
                    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    translated_doc.add_paragraph("")

                translated_doc_filename = "translated_document.docx"
                translated_doc.save(translated_doc_filename)

            # --- PDF Output ---
            with st.spinner("Generating PDF file..."):
                translated_pdf_filename = "translated_document.pdf"
                create_pdf_from_paragraphs(translated_paragraphs, translated_pdf_filename, is_rtl=is_target_rtl)

            # --- Download Button (ZIP) ---
            st.subheader("Download Files")
            
            # Create a BytesIO object to hold the zip file in memory
            zip_buffer = io.BytesIO()
            
            with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
                zip_file.write(combined_doc_filename, arcname=combined_doc_filename)
                zip_file.write(translated_doc_filename, arcname=translated_doc_filename)
                zip_file.write(translated_pdf_filename, arcname=translated_pdf_filename)
            
            st.download_button(
                label="Download All Files (ZIP)",
                data=zip_buffer.getvalue(),
                file_name="translated_files.zip",
                mime="application/zip",
            )
            
            # Clean up temporary files (optional, but good practice)
            os.remove(combined_doc_filename)
            os.remove(translated_doc_filename)
            os.remove(translated_pdf_filename)

if __name__ == "__main__":
    main()
